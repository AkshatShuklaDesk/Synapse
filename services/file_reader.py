"""
File reading utilities for PDF and DOCX resume files.
"""

from pathlib import Path
from pypdf import PdfReader
from docx import Document


def read_pdf(file_path: str | Path) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx(file_path: str | Path) -> str:
    """Extract text from a DOCX file, including table content."""
    document = Document(file_path)
    text = ""

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def read_resume_file(file_path: str | Path) -> str | None:
    """Read a resume file and return its text content.

    Supports .pdf and .docx formats.
    Returns None if the format is unsupported.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return read_pdf(file_path)
    elif suffix == ".docx":
        return read_docx(file_path)
    else:
        return None
